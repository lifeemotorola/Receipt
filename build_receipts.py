#!/usr/bin/env python3
"""
Suahco4 — Receipt Book builder.

Produces:
  receipt_book.pdf   — print-ready, 2 receipts per A4 page (Original + Carbon Copy)
  receipt_book.docx  — same book, fully editable in Word

Edit the CONFIG block below to customise anything, then run:  python3 build_receipts.py
"""

# ============================== CONFIG ==============================
CONFIG = {
    # --- header ---
    "biz":      "Suahco4",
    "addr":     "COW FARM, NEW ISRAEL COMMUNITY, BREWERVILLE -LIBERIA",
    "contact":  "Phone: 0778662590 | Email:",
    "extra":    "",
    "title":    "RECEIPT",
    "logo":     "logo.png",        # path, or "" for no logo
    "logo_w_mm": 14,

    # --- numbering ---
    "prefix":   "SUAHCO4-",
    "suffix":   "",
    "start":    61,
    "pad":      3,
    "count":    50,                # how many receipt numbers
    "step":     1,
    "no_label": "Receipt No:",
    "seq_tpl":  "[Receipt #{n}]",  # "" to hide

    # --- fields:  ("line", label, extra_blank_lines, pair_with_next, prefilled_value)
    #              ("opts", label, "Opt, Opt, Opt",   pair_with_next, "")
    "fields": [
        ("line", "Date:",          0, False, ""),
        ("line", "Received from:", 0, False, ""),
        ("line", "The sum of:",    1, False, ""),
        ("line", "For:",           1, False, ""),
        ("line", "Amount Paid$:",  0, True,  ""),
        ("line", "Balance:",       0, False, ""),
        ("opts", "Payment Method:", "Cash, Check, Card, Transfer", False, ""),
        ("opts", "Currency:",       "USD, LRD",                    False, ""),
    ],

    # --- footer ---
    "sig_cap":   "Authorized Signature",
    "foot_note": "",
    "show_sig":  True,

    # --- copies ---
    "copies":       2,
    "copy_labels":  ["ORIGINAL", "CARBON COPY"],
    "copy_tints":   ["#ffffff", "#d9eef7"],   # background of each copy
    "show_copy_tag": True,
    # how the copies are laid out across pages:
    #   "page"     = a page of ORIGINALS, then a page of the SAME numbers as copies  (best: stack & cut together)
    #   "together" = each number's copies sit back-to-back in the flow
    #   "separate" = every original first, then every copy
    "copy_order": "page",

    # --- page ---
    "page_w_mm": 210, "page_h_mm": 297,
    "per_page":  3,
    "margin_mm": 7,
    "gutter_mm": 3,        # gap between receipts on a page
    "cut_line":  "dash",   # dash | solid | none
    "cut_all":   True,     # also draw a cut line under the LAST receipt on the page
    "border":    True,

    # --- type ---
    "font":       "Helvetica",       # Helvetica | Times | Courier
    "fs":         7.5,
    "fs_biz":     13,
    "fs_addr":    5.8,
    "fs_title":   8.5,
    "accent":     "#B8121B",
    "text":       "#111111",
    "row_gap_mm": 1.1,
    "label_col_mm": 0,      # 0 = auto: all writing lines start in one straight column
    "align_fields": "fill", # fill | center  (how rows use the free height)

    # ---- TEMPLATE ----------------------------------------------------
    # "simple"   = the blank-lines receipt (your original book)
    # "itemized" = vertical table of items with qty / rate / amount + totals
    "template": "simple",

    # header arrangement: "stacked" = centred block (portrait look)
    #                     "inline"  = name on the left, title + number on the right (horizontal look)
    "header_style": "stacked",

    # ---- itemized-template settings ----------------------------------
    # each column: (heading, width weight, alignment)   align = l | c | r
    "item_cols": [
        ("#",            0.6, "c"),
        ("DESCRIPTION",  6.0, "l"),
        ("QTY",          1.0, "c"),
        ("RATE",         1.6, "r"),
        ("AMOUNT",       1.9, "r"),
    ],
    "item_rows":      6,                # blank writing rows in the table
    "totals":         ["Sub-Total", "Discount", "TOTAL"],
    "head_fill":      "#e9edf2",        # table heading shading
    "total_fill":     "#f4f6f8",
    "grid":           "#333333",
    # rows printed above the table on the itemized receipt
    "item_top": [
        ("line", "Date:",          0, True,  ""),
        ("line", "Received from:", 0, False, ""),
    ],
    # ---- fees-template settings --------------------------------------
    "fee_cols": [
        ("PARTICULARS", 4.6, "l"),
        ("AMOUNT DUE",  2.0, "r"),
        ("AMOUNT PAID", 2.0, "r"),
        ("BALANCE",     2.0, "r"),
    ],
    "fee_rows": [
        "Tuition Fee",
        "Registration Fee",
        "Uniform",
        "Books / Materials",
        "Examination Fee",
        "Other (specify)",
    ],
    "fee_total": "TOTAL",
    "fee_top": [
        ("line", "Student Name:", 0, False, ""),
        ("line", "Student ID:",   0, True,  ""),
        ("line", "Class / Grade:", 0, False, ""),
        ("line", "Academic Year:", 0, True,  ""),
        ("opts", "Term:", "1st Term, 2nd Term, 3rd Term", False, ""),
    ],
    "fee_bottom": [
        ("line", "Amount in words:", 0, False, ""),
        ("opts", "Payment Method:", "Cash, Check, Card, Transfer", False, ""),
        ("opts", "Currency:", "USD, LRD", False, ""),
        ("line", "Balance Carried Forward:", 0, True, ""),
        ("line", "Next Payment Due:", 0, False, ""),
    ],
    "sig_caps": [],          # e.g. ["Parent / Guardian", "Bursar / Cashier"] for two lines

    # rows printed below the table
    "item_bottom": [
        ("line", "Amount in words:", 0, False, ""),
        ("opts", "Payment Method:", "Cash, Check, Card, Transfer", False, ""),
        ("opts", "Currency:",       "USD, LRD",                    False, ""),
    ],
}
# ====================================================================

import os
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor, black
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

C = CONFIG
BOLD = {"Helvetica": "Helvetica-Bold", "Times": "Times-Bold", "Courier": "Courier-Bold"}
REG = {"Helvetica": "Helvetica", "Times": "Times-Roman", "Courier": "Courier"}
FB, FR = BOLD[C["font"]], REG[C["font"]]
ACC, TXT = HexColor(C["accent"]), HexColor(C["text"])


def numbers():
    return [C["prefix"] + str(C["start"] + i * C["step"]).zfill(C["pad"]) + C["suffix"]
            for i in range(C["count"])]


def sequence():
    """[(receipt_index, copy_index), ...] in printing order."""
    n, cp, per = C["count"], max(1, C["copies"]), max(1, C["per_page"])
    mode = C.get("copy_order", "page")
    if mode == "separate":
        return [(i, c) for c in range(cp) for i in range(n)]
    if mode == "together":
        return [(i, c) for i in range(n) for c in range(cp)]
    # "page": fill a page with originals, then repeat that same page for each copy
    out = []
    for blk in range(0, n, per):
        group = list(range(blk, min(blk + per, n)))
        for c in range(cp):
            out += [(i, c) for i in group]
            if len(group) < per:          # pad the short last page so copies stay aligned
                out += [(None, c)] * (per - len(group))
    return out


# ------------------------------------------------------------------ PDF
def _rows_from(fields):
    """Turn a field list into laid-out rows (handles pairing + ruled lines)."""
    rows, i = [], 0
    while i < len(fields):
        f = fields[i]
        if f[0] == "line":
            _, lab, ruled, pair, val = f
            if pair and i + 1 < len(fields) and ruled == 0:
                g = fields[i + 1]
                second = ("line", g[1], 0, g[4]) if g[0] == "line" else ("opts", g[1], g[2], g[4])
                rows.append([("line", lab, 0, val), second])
                i += 2
                continue
            rows.append([("line", lab, 0, val)])
            for _k in range(ruled):
                rows.append([("cont", "", 0, "")])
            i += 1
        else:
            _, lab, opts, pair, val = f
            if pair and i + 1 < len(fields):
                g = fields[i + 1]
                second = ("line", g[1], 0, g[4]) if g[0] == "line" else ("opts", g[1], g[2], g[4])
                rows.append([("opts", lab, opts, val), second])
                i += 2
                continue
            rows.append([("opts", lab, opts, val)])
            i += 1
    return rows


def _label_col(cv, rows):
    if C["label_col_mm"] > 0:
        return C["label_col_mm"] * mm
    widest = 0
    for row in rows:
        for kind, lab, _p, _v in row:
            if kind != "cont":
                widest = max(widest, cv.stringWidth(lab, FB, C["fs"]))
    return widest + 2.2 * mm


def _draw_rows(cv, rows, L, R, top, rh, lab_col):
    """Draw laid-out rows downward from `top`; returns the y after the last row."""
    ry = top
    for row in rows:
        ry -= rh
        cells = len(row)
        col_w = (R - L) / cells
        gut = 4 * mm
        for j, (kind, lab, payload, val) in enumerate(row):
            cxl = L + j * col_w
            cxr = (cxl + col_w - gut) if j < cells - 1 else R
            lc = min(lab_col, col_w * 0.5)
            cv.setFillColor(TXT)
            if kind == "cont":
                cv.setStrokeColor(black); cv.setLineWidth(0.7)
                cv.line(cxl + lc, ry, cxr, ry)
                continue
            cv.setFont(FB, C["fs"])
            cv.drawString(cxl, ry + 1.1, lab)
            if kind == "line":
                cv.setStrokeColor(black); cv.setLineWidth(0.7)
                cv.line(cxl + lc, ry, cxr, ry)
                if val:
                    cv.setFont(FR, C["fs"])
                    cv.drawString(cxl + lc + 1 * mm, ry + 1.4, val)
            else:
                opts = [o.strip() for o in payload.split(",") if o.strip()]
                bs = C["fs"] * 1.1
                gapo = 3.4 * mm
                total = sum(bs + 1.3 * mm + cv.stringWidth(o, FR, C["fs"]) for o in opts) + gapo * (len(opts) - 1)
                nat = [bs + 1.3 * mm + cv.stringWidth(o, FR, C["fs"]) for o in opts]
                pitch = None
                if total < (cxr - cxl - lc):
                    pitch = min((cxr - cxl - lc) / len(opts), max(nat) + gapo * 1.4)
                bx = cxl + lc
                cv.setFont(FR, C["fs"])
                for k, opt in enumerate(opts):
                    cv.setStrokeColor(black); cv.setLineWidth(0.7)
                    cv.rect(bx, ry, bs, bs, stroke=1, fill=0)
                    cv.setFillColor(TXT)
                    cv.drawString(bx + bs + 1.3 * mm, ry + 1.4, opt)
                    bx = (cxl + lc + (k + 1) * pitch) if pitch else \
                         bx + bs + 1.3 * mm + cv.stringWidth(opt, FR, C["fs"]) + gapo
    return ry


def draw_items_table(cv, L, R, top, bottom):
    """Vertical item table between `top` and `bottom`; returns y below the table."""
    cols = C["item_cols"]
    tw = sum(c[1] for c in cols)
    widths = [(c[1] / tw) * (R - L) for c in cols]
    xs, acc = [], L
    for wd in widths:
        xs.append(acc); acc += wd
    xs.append(R)

    n_body = C["item_rows"]
    n_tot = len(C["totals"])
    head_h = C["fs"] * 1.9
    avail = top - bottom
    row_h = min((avail - head_h) / (n_body + n_tot), 7.2 * mm)
    row_h = max(row_h, C["fs"] * 1.45)

    grid = HexColor(C["grid"])
    y = top

    # heading band
    cv.setFillColor(HexColor(C["head_fill"]))
    cv.rect(L, y - head_h, R - L, head_h, stroke=0, fill=1)
    cv.setFont(FB, C["fs"] * 0.95)
    cv.setFillColor(TXT)
    for i, (hd, _w, al) in enumerate(cols):
        ty = y - head_h + head_h * 0.32
        if al == "l":   cv.drawString(xs[i] + 1.6 * mm, ty, hd)
        elif al == "r": cv.drawRightString(xs[i + 1] - 1.6 * mm, ty, hd)
        else:           cv.drawCentredString((xs[i] + xs[i + 1]) / 2, ty, hd)
    y -= head_h

    # body rows
    cv.setFont(FR, C["fs"])
    for r in range(n_body):
        cv.setFillColor(TXT)
        cv.drawCentredString((xs[0] + xs[1]) / 2, y - row_h + row_h * 0.32, str(r + 1))
        y -= row_h
        cv.setStrokeColor(grid); cv.setLineWidth(0.5)
        cv.line(L, y, R, y)
    body_bot = y

    # totals block — label spans the columns left of the last one
    lab_x = xs[-3] if len(xs) >= 3 else xs[0]
    for i, t in enumerate(C["totals"]):
        last = (i == len(C["totals"]) - 1)
        cv.setFillColor(HexColor(C["total_fill"]))
        cv.rect(lab_x, y - row_h, R - lab_x, row_h, stroke=0, fill=1)
        cv.setFillColor(TXT)
        cv.setFont(FB if last else FR, C["fs"] * (1.02 if last else 0.95))
        cv.drawRightString(xs[-2] - 1.6 * mm, y - row_h + row_h * 0.32, t)
        y -= row_h
        cv.setStrokeColor(grid); cv.setLineWidth(1.0 if last else 0.5)
        cv.line(lab_x, y, R, y)
    tot_bot = y

    # grid frame + verticals
    cv.setStrokeColor(grid); cv.setLineWidth(0.9)
    cv.rect(L, body_bot, R - L, top - body_bot, stroke=1, fill=0)
    cv.setLineWidth(0.5)
    cv.line(L, top - head_h, R, top - head_h)
    for i in range(1, len(cols)):
        cv.line(xs[i], body_bot, xs[i], top)
        if xs[i] >= lab_x:
            cv.line(xs[i], tot_bot, xs[i], body_bot)
    cv.setLineWidth(0.9)
    cv.line(lab_x, tot_bot, lab_x, body_bot)
    cv.line(R, tot_bot, R, body_bot)
    return tot_bot


def draw_fee_table(cv, L, R, top, bottom):
    """Named fee-breakdown rows + a TOTAL band. Returns y below the table."""
    cols = C["fee_cols"]
    tw = sum(c[1] for c in cols)
    widths = [(c[1] / tw) * (R - L) for c in cols]
    xs, acc = [], L
    for wd in widths:
        xs.append(acc); acc += wd
    xs.append(R)

    labels = C["fee_rows"]
    head_h = C["fs"] * 1.9
    n = len(labels) + 1                       # + total row
    row_h = min((top - bottom - head_h) / n, 7.0 * mm)
    row_h = max(row_h, C["fs"] * 1.5)
    grid = HexColor(C["grid"])
    y = top

    cv.setFillColor(HexColor(C["head_fill"]))
    cv.rect(L, y - head_h, R - L, head_h, stroke=0, fill=1)
    cv.setFont(FB, C["fs"] * 0.95)
    cv.setFillColor(TXT)
    for i, (hd, _w, al) in enumerate(cols):
        ty = y - head_h + head_h * 0.32
        if al == "l":   cv.drawString(xs[i] + 1.6 * mm, ty, hd)
        elif al == "r": cv.drawRightString(xs[i + 1] - 1.6 * mm, ty, hd)
        else:           cv.drawCentredString((xs[i] + xs[i + 1]) / 2, ty, hd)
    y -= head_h

    cv.setFont(FR, C["fs"])
    for lab in labels:
        cv.setFillColor(TXT)
        cv.drawString(xs[0] + 1.6 * mm, y - row_h + row_h * 0.32, lab)
        y -= row_h
        cv.setStrokeColor(grid); cv.setLineWidth(0.5)
        cv.line(L, y, R, y)

    cv.setFillColor(HexColor(C["total_fill"]))
    cv.rect(L, y - row_h, R - L, row_h, stroke=0, fill=1)
    cv.setFillColor(TXT)
    cv.setFont(FB, C["fs"] * 1.05)
    cv.drawRightString(xs[1] - 1.6 * mm, y - row_h + row_h * 0.32, C["fee_total"])
    y -= row_h

    cv.setStrokeColor(grid); cv.setLineWidth(0.9)
    cv.rect(L, y, R - L, top - y, stroke=1, fill=0)
    cv.setLineWidth(0.5)
    cv.line(L, top - head_h, R, top - head_h)
    cv.setLineWidth(0.9)
    cv.line(L, y + row_h, R, y + row_h)
    cv.setLineWidth(0.5)
    for i in range(1, len(cols)):
        cv.line(xs[i], y, xs[i], top)
    return y


def draw_receipt(cv, x, y, w, h, idx, copy):
    """x,y = bottom-left of the receipt cell (points)."""
    label = C["copy_labels"][copy] if copy < len(C["copy_labels"]) else ""
    tint = C["copy_tints"][copy] if copy < len(C["copy_tints"]) else "#ffffff"
    no = numbers()[idx]
    seq_n = C["start"] + idx * C["step"]

    if tint.lower() not in ("#ffffff", "#fff"):
        cv.setFillColor(HexColor(tint))
        cv.rect(x, y, w, h, stroke=0, fill=1)

    pad = 4 * mm
    if C["border"]:
        cv.setStrokeColor(black)
        cv.setLineWidth(1)
        cv.rect(x + 2 * mm, y + 2 * mm, w - 4 * mm, h - 4 * mm, stroke=1, fill=0)

    L = x + pad + 2 * mm
    R = x + w - pad - 2 * mm
    top = y + h - pad - 2 * mm
    cur = top

    # ---- logo
    logo_w = 0
    if C["logo"] and os.path.exists(C["logo"]):
        img = ImageReader(C["logo"])
        iw, ih = img.getSize()
        logo_w = C["logo_w_mm"] * mm
        logo_h = logo_w * ih / iw
        cv.drawImage(img, L, cur - logo_h, logo_w, logo_h, mask="auto")

    # ---- header ------------------------------------------------------
    if C["header_style"] == "inline":
        hx = L + (logo_w + 3 * mm if logo_w else 0)
        cv.setFillColor(TXT)
        cv.setFont(FB, C["fs_biz"])
        cur -= C["fs_biz"] * 0.95
        cv.drawString(hx, cur, C["biz"])
        # title + number stack on the right
        ty = cur
        cv.setFont(FB, C["fs_title"])
        cv.drawRightString(R, ty, C["title"])
        sub = cur
        for line in (C["addr"], C["contact"], C["extra"]):
            if line:
                sub -= C["fs_addr"] * 1.45
                cv.setFont(FR, C["fs_addr"])
                cv.setFillColor(TXT)
                cv.drawString(hx, sub, line)
        ny = ty - C["fs"] * 1.9
        cv.setFont(FB, C["fs"] * 1.1)
        cv.setFillColor(ACC)
        cv.drawRightString(R, ny, f'{C["no_label"]} {no}')
        if C["show_copy_tag"] and label:
            cv.setFillColor(TXT)
            cv.setFont(FB, C["fs"] * 0.92)
            cv.drawRightString(R, ny - C["fs"] * 1.5, label)
        cur = min(sub, ny - C["fs"] * 1.5) - C["fs"] * 0.6
        cv.setStrokeColor(black)
        cv.setLineWidth(0.9)
        cv.line(L, cur, R, cur)
        cur -= C["fs"] * 0.4
    else:
        cx = (L + R) / 2
        cv.setFillColor(TXT)
        cv.setFont(FB, C["fs_biz"])
        cur -= C["fs_biz"] * 0.95
        cv.drawCentredString(cx, cur, C["biz"])
        for line in (C["addr"], C["contact"], C["extra"]):
            if line:
                cur -= C["fs_addr"] * 1.45
                cv.setFont(FR, C["fs_addr"])
                cv.drawCentredString(cx, cur, line)
        cur -= C["fs_title"] * 1.7
        cv.setFont(FB, C["fs_title"])
        cv.drawCentredString(cx, cur, C["title"])

        # ---- meta row
        cur -= C["fs"] * 2.0
        cv.setFont(FB, C["fs"] * 1.1)
        cv.setFillColor(ACC)
        cv.drawString(L, cur, f'{C["no_label"]} {no}')
        if C["show_copy_tag"] and label:
            cv.setFillColor(TXT)
            cv.setFont(FB, C["fs"] * 0.92)
            cv.drawRightString(R, cur, label)

    # ---- footer geometry (reserve space)
    foot_y = y + pad + 2 * mm
    caps = C.get("sig_caps") or ([C["sig_cap"]] if C["show_sig"] else [])
    sig_w = 42 * mm
    if caps:
        cv.setFont(FR, C["fs"] * 0.85)
        if len(caps) == 1:
            spots = [(R - sig_w, R)]
        else:                                  # spread across the width, right-anchored
            sig_w = min(48 * mm, (R - L) / len(caps) - 6 * mm)
            spots = []
            for k in range(len(caps)):
                cx0 = L + k * ((R - L) / len(caps))
                spots.append((cx0, cx0 + sig_w))
        for (a, b), cap in zip(spots, caps):
            cv.setStrokeColor(black); cv.setLineWidth(0.7)
            cv.line(a, foot_y + C["fs"] * 1.3, b, foot_y + C["fs"] * 1.3)
            cv.setFillColor(TXT)
            cv.drawCentredString((a + b) / 2, foot_y + C["fs"] * 0.25, cap)
    fy = foot_y + C["fs"] * 0.25
    seq_txt = C["seq_tpl"].replace("{n}", str(seq_n)) if C["seq_tpl"] else ""
    cv.setFont(FR, C["fs"] * 0.8)
    cv.setFillColor(HexColor("#555555"))
    if len(caps) >= 2:                      # sit below the signature captions
        by = foot_y - C["fs"] * 1.05
        if seq_txt:
            cv.drawString(L, by, seq_txt)
        if C["foot_note"]:
            cv.drawCentredString((L + R) / 2, by, C["foot_note"])
    else:
        if C["foot_note"]:
            cv.setFont(FR, C["fs"] * 0.85)
            cv.drawString(L, fy + C["fs"] * 1.3, C["foot_note"])
        if seq_txt:
            cv.setFont(FR, C["fs"] * 0.85)
            cv.drawString(L, fy, seq_txt)

    # ---- body -------------------------------------------------------
    if C["template"] == "fees":
        top_rows = _rows_from(C["fee_top"])
        bot_rows = _rows_from(C["fee_bottom"])
        lab_col = _label_col(cv, top_rows + bot_rows)
        rh = C["fs"] * 2.15
        after_top = _draw_rows(cv, top_rows, L, R, cur - C["fs"] * 1.7, rh, lab_col)
        bot_need = len(bot_rows) * rh + C["fs"] * 1.0
        tbl_bot = foot_y + C["fs"] * (3.6 if len(C.get("sig_caps") or []) >= 2 else 2.8) + bot_need
        after_tbl = draw_fee_table(cv, L, R, after_top - C["fs"] * 1.2, tbl_bot)
        _draw_rows(cv, bot_rows, L, R, after_tbl - C["fs"] * 1.0, rh, lab_col)
        return

    if C["template"] == "itemized":
        top_rows = _rows_from(C["item_top"])
        bot_rows = _rows_from(C["item_bottom"])
        lab_col = _label_col(cv, top_rows + bot_rows)
        rh = C["fs"] * 2.2
        zone_top = cur - C["fs"] * 1.7
        after_top = _draw_rows(cv, top_rows, L, R, zone_top, rh, lab_col)
        bot_need = len(bot_rows) * rh + C["fs"] * 1.2
        tbl_bot = foot_y + C["fs"] * 2.6 + bot_need
        after_tbl = draw_items_table(cv, L, R, after_top - C["fs"] * 1.1, tbl_bot)
        _draw_rows(cv, bot_rows, L, R, after_tbl - C["fs"] * 0.9, rh, lab_col)
        return

    rows = _rows_from(C["fields"])
    zone_top = cur - C["fs"] * 1.9
    zone_bot = foot_y + C["fs"] * 3.0
    n_rows = max(1, len(rows))
    avail = zone_top - zone_bot
    rh = min(avail / n_rows, 8.6 * mm)
    lab_col = _label_col(cv, rows)
    if C["align_fields"] == "center":
        zone_top -= (avail - rh * n_rows) / 2
    _draw_rows(cv, rows, L, R, zone_top, rh, lab_col)


def build_pdf(path="receipt_book.pdf"):
    W, H = C["page_w_mm"] * mm, C["page_h_mm"] * mm
    cv = canvas.Canvas(path, pagesize=(W, H))
    cv.setTitle(f'{C["biz"]} — Receipt Book')
    m, g, per = C["margin_mm"] * mm, C["gutter_mm"] * mm, C["per_page"]
    cell_h = (H - 2 * m - g * (per - 1)) / per
    cell_w = W - 2 * m

    seq = sequence()
    for p in range(0, len(seq), per):
        chunk = seq[p:p + per]
        for k, (idx, copy) in enumerate(chunk):
            y = H - m - (k + 1) * cell_h - k * g
            if idx is not None:
                draw_receipt(cv, m, y, cell_w, cell_h, idx, copy)
            if (k < per - 1 or C.get("cut_all")) and C["cut_line"] != "none":
                ly = y - g / 2 if k < per - 1 else y - g / 2
                cv.setStrokeColor(HexColor("#888888"))
                cv.setLineWidth(0.8)
                if C["cut_line"] == "dash":
                    cv.setDash(3, 3)
                cv.line(m, ly, W - m, ly)
                cv.setDash()
                cv.setFont(FR, 6)
                cv.setFillColor(HexColor("#888888"))
                cv.drawString(m, ly + 1.2, "\u2702")
        cv.showPage()
    cv.save()
    return path, len(seq)


# ------------------------------------------------------------------ DOCX
def build_docx(path="receipt_book.docx"):
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(C["page_w_mm"]), Mm(C["page_h_mm"])
    for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, a, Mm(C["margin_mm"]))
    st = doc.styles["Normal"]
    st.font.name = "Arial" if C["font"] == "Helvetica" else C["font"]
    st.font.size = Pt(C["fs"])
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.space_before = Pt(0)

    def shade(cell, hexcolor):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexcolor.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(el)

    def borders(tbl, on=True):
        tblPr = tbl._tbl.tblPr
        b = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single" if on else "none")
            e.set(qn("w:sz"), "8")
            e.set(qn("w:color"), "222222")
            b.append(e)
        for edge in ("insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            b.append(e)
        tblPr.append(b)

    def para(cell, text, size, bold=False, align="left", color=None, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                       "center": WD_ALIGN_PARAGRAPH.CENTER,
                       "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color.lstrip("#").upper())
        return p

    UL = "_"
    nums = numbers()
    seq = sequence()

    for pos, (idx, copy) in enumerate(seq):
        if idx is None:
            continue
        label = C["copy_labels"][copy] if copy < len(C["copy_labels"]) else ""
        tint = C["copy_tints"][copy] if copy < len(C["copy_tints"]) else "#ffffff"

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        borders(tbl, C["border"])
        cell = tbl.cell(0, 0)
        shade(cell, tint)

        # header
        if C["logo"] and os.path.exists(C["logo"]):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run().add_picture(C["logo"], width=Mm(C["logo_w_mm"]))
            para(cell, C["biz"], C["fs_biz"], True, "center")
        else:
            para(cell, C["biz"], C["fs_biz"], True, "center", first=True)
        for line in (C["addr"], C["contact"], C["extra"]):
            if line:
                para(cell, line, C["fs_addr"], False, "center")
        para(cell, C["title"], C["fs_title"], True, "center")

        p = cell.add_paragraph()
        r = p.add_run(f'{C["no_label"]} {nums[idx]}')
        r.bold = True
        r.font.size = Pt(C["fs"] * 1.1)
        r.font.color.rgb = RGBColor.from_string(C["accent"].lstrip("#").upper())
        if C["show_copy_tag"] and label:
            r2 = p.add_run("\t" + label)
            r2.bold = True
            r2.font.size = Pt(C["fs"] * 0.92)

        def simple_rows(flds, cell):
            i = 0
            while i < len(flds):
                f = flds[i]
                if f[0] == "line":
                    _, lab, ruled, pair, val = f
                    if pair and i + 1 < len(flds) and ruled == 0:
                        g = flds[i + 1]
                        para(cell, f"{lab} {val or UL*18}    {g[1]} {g[4] or UL*18}", C["fs"])
                        i += 2
                        continue
                    para(cell, f"{lab} {val or UL*(56-len(lab))}", C["fs"])
                    i += 1
                else:
                    _, lab, opts, pair, val = f
                    boxes = "   ".join("\u2610 " + o.strip() for o in opts.split(",") if o.strip())
                    para(cell, f"{lab}  {boxes}", C["fs"])
                    i += 1

        def close_slip(cell, idx, pos):
            cell.add_paragraph()
            seq_txt = C["seq_tpl"].replace("{n}", str(C["start"] + idx * C["step"])) if C["seq_tpl"] else ""
            caps = C.get("sig_caps") or ([C["sig_cap"]] if C["show_sig"] else [])
            if len(caps) >= 2:
                para(cell, "".join(UL*22 + "\t\t" for _ in caps), C["fs"] * 0.85)
                para(cell, "".join(c.ljust(24) + "\t\t" for c in caps), C["fs"] * 0.85)
                if seq_txt:
                    para(cell, seq_txt, C["fs"] * 0.85)
            elif caps:
                para(cell, f'{seq_txt}\t\t\t\t{UL*26}', C["fs"] * 0.85)
                para(cell, "\t\t\t\t" + caps[0], C["fs"] * 0.85)
            elif seq_txt:
                para(cell, seq_txt, C["fs"] * 0.85)
            if pos < len(seq) - 1:
                sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if (pos % C["per_page"]) == C["per_page"] - 1:
                    sep.add_run().add_break(WD_BREAK.PAGE)
                else:
                    r = sep.add_run("\u2702 " + "- " * 45)
                    r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # ---- fees body ---------------------------------------------
        if C["template"] == "fees":
            simple_rows(C["fee_top"], cell)
            cell.add_paragraph()
            cols = C["fee_cols"]
            labels = C["fee_rows"]
            ft = cell.add_table(rows=1 + len(labels) + 1, cols=len(cols))
            ft.style = "Table Grid"
            tw = sum(c[1] for c in cols)
            usable = C["page_w_mm"] - 2 * C["margin_mm"] - 16
            for ci, (hd, wgt, al) in enumerate(cols):
                hc = ft.cell(0, ci)
                shade(hc, C["head_fill"])
                para(hc, hd, C["fs"] * 0.95, True,
                     {"l": "left", "c": "center", "r": "right"}[al], first=True)
                for ri in range(1 + len(labels) + 1):
                    ft.cell(ri, ci).width = Mm(usable * wgt / tw)
            for ri, lab in enumerate(labels):
                para(ft.cell(1 + ri, 0), lab, C["fs"], False, "left", first=True)
            tr = 1 + len(labels)
            for ci in range(len(cols)):
                shade(ft.cell(tr, ci), C["total_fill"])
            para(ft.cell(tr, 0), C["fee_total"], C["fs"] * 1.05, True, "right", first=True)
            cell.add_paragraph()
            simple_rows(C["fee_bottom"], cell)
            close_slip(cell, idx, pos)
            continue

        # ---- itemized body -----------------------------------------
        if C["template"] == "itemized":
            simple_rows(C["item_top"], cell)
            cell.add_paragraph()

            cols = C["item_cols"]
            it = cell.add_table(rows=1 + C["item_rows"] + len(C["totals"]), cols=len(cols))
            it.style = "Table Grid"
            tw = sum(c[1] for c in cols)
            usable = C["page_w_mm"] - 2 * C["margin_mm"] - 16
            for ci, (hd, wgt, al) in enumerate(cols):
                cw = Mm(usable * wgt / tw)
                hc = it.cell(0, ci)
                shade(hc, C["head_fill"])
                para(hc, hd, C["fs"] * 0.95, True,
                     {"l": "left", "c": "center", "r": "right"}[al], first=True)
                for ri in range(1 + C["item_rows"] + len(C["totals"])):
                    it.cell(ri, ci).width = cw
            for r in range(C["item_rows"]):
                para(it.cell(1 + r, 0), str(r + 1), C["fs"], False, "center", first=True)
            base = 1 + C["item_rows"]
            for ti, t in enumerate(C["totals"]):
                last = ti == len(C["totals"]) - 1
                lc = it.cell(base + ti, len(cols) - 2)
                merged = it.cell(base + ti, 0).merge(lc)
                shade(merged, C["total_fill"])
                shade(it.cell(base + ti, len(cols) - 1), C["total_fill"])
                para(merged, t, C["fs"] * (1.02 if last else 0.95), last, "right", first=True)

            cell.add_paragraph()
            simple_rows(C["item_bottom"], cell)
            close_slip(cell, idx, pos)
            continue

        # fields
        i, flds = 0, C["fields"]
        while i < len(flds):
            f = flds[i]
            if f[0] == "line":
                _, lab, ruled, pair, val = f
                if pair and i + 1 < len(flds) and ruled == 0:
                    g = flds[i + 1]
                    para(cell, f"{lab} {val or UL*22}    {g[1]} {g[4] or UL*22}", C["fs"])
                    i += 2
                    continue
                para(cell, f"{lab} {val or UL*(58-len(lab))}", C["fs"])
                for _k in range(ruled):
                    para(cell, UL * 62, C["fs"])
                i += 1
            else:
                _, lab, opts, pair, val = f
                boxes = "   ".join("\u2610 " + o.strip() for o in opts.split(",") if o.strip())
                para(cell, f"{lab}  {boxes}", C["fs"])
                i += 1

        cell.add_paragraph()
        if C["foot_note"]:
            para(cell, C["foot_note"], C["fs"] * 0.85)
        seq_txt = C["seq_tpl"].replace("{n}", str(C["start"] + idx * C["step"])) if C["seq_tpl"] else ""
        if C["show_sig"]:
            para(cell, f'{seq_txt}\t\t\t\t{UL*26}', C["fs"] * 0.85, False, "left")
            para(cell, "\t\t\t\t" + C["sig_cap"], C["fs"] * 0.85, False, "left")
        elif seq_txt:
            para(cell, seq_txt, C["fs"] * 0.85)

        last_on_page = (pos % C["per_page"]) == C["per_page"] - 1
        if pos < len(seq) - 1:
            sep = doc.add_paragraph()
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if last_on_page:
                run = sep.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                r = sep.add_run("\u2702 " + "- " * 45)
                r.font.size = Pt(7)
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(path)
    return path, len(seq)


# ------------------------------------------------------------------ books
ITEMIZED = {                 # overrides applied on top of CONFIG
    "template": "itemized",
    "title": "SALES RECEIPT",
    "per_page": 2,
    "fs": 8, "fs_biz": 15, "fs_addr": 6.2, "fs_title": 9.5,
}


FEES = {
    "template": "fees",
    "title": "SCHOOL FEES PAYMENT SLIP",
    "per_page": 2,
    "fs": 8, "fs_biz": 15, "fs_addr": 6.2, "fs_title": 9.5,
    "sig_caps": ["Parent / Guardian", "Bursar / Cashier"],
    "foot_note": "Keep this slip safe \u2014 it is your proof of payment.",
}


WIDE = {                     # horizontal / landscape slip
    "template": "simple",
    "header_style": "inline",
    "title": "PAYMENT RECEIPT",
    "page_w_mm": 297, "page_h_mm": 210,     # A4 landscape
    "per_page": 2,
    "margin_mm": 8, "gutter_mm": 4,
    "fs": 8.5, "fs_biz": 15, "fs_addr": 6.4, "fs_title": 11,
    "fields": [
        ("line", "Date:",           0, True,  ""),
        ("line", "Received from:",  0, False, ""),
        ("line", "The sum of:",     1, False, ""),
        ("line", "Being payment for:", 0, False, ""),
        ("line", "Amount Paid$:",   0, True,  ""),
        ("line", "Balance:",        0, False, ""),
        ("opts", "Payment Method:", "Cash, Check, Card, Transfer", True, ""),
        ("opts", "Currency:",       "USD, LRD",                    False, ""),
    ],
}


def build(name, overrides=None):
    global C
    base = dict(CONFIG)
    base.update(overrides or {})
    C = base
    build_pdf(f"{name}.pdf")
    build_docx(f"{name}.docx")
    return f"{name}.pdf / {name}.docx"


if __name__ == "__main__":
    print("Template 1 (blank lines)  ->", build("receipt_book"))
    print("Template 2 (itemized)     ->", build("receipt_book_itemized", ITEMIZED))
    print("Template 3 (school fees)  ->", build("receipt_book_fees", FEES))
    print("Template 4 (horizontal)   ->", build("receipt_book_wide", WIDE))
